import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openai import OpenAI

def create_cover_page(doc):
    # Add logo (replace with your image path if needed)
    try:
        doc.add_picture('logo.png', width=Inches(2))
    except:
        pass  # Skip if logo not found

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("TENDER ANALYSIS REPORT")
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 32, 96)  # Dark blue
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'LightShading-Accent1'
    
    cells = [
        ("Tender Reference:", "B2-2024"),
        ("Client:", "Ministry of Infrastructure"),
        ("Submission Deadline:", datetime.now().strftime("%d %B %Y")),
        ("Prepared By:", "Your Company Name")
    ]
    
    for i, (label, value) in enumerate(cells):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        table.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_page_break()

def add_methodology_table(doc):
    doc.add_heading('Implementation Methodology', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'MediumGrid3-Accent1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Activities'
    hdr_cells[2].text = 'Timeline'
    
    # Sample data (replace with AI-generated content)
    phases = [
        ('1. Preparation', 'Site assessment\nTeam mobilization', 'Weeks 1-2'),
        ('2. Construction', 'Raw water intake installation\nQuality testing', 'Weeks 3-8'),
        ('3. Handover', 'Final inspection\nDocumentation', 'Week 9')
    ]
    
    for phase, activities, timeline in phases:
        row_cells = table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = activities
        row_cells[2].text = timeline

def generate_reports():
    client = OpenAI()
    
    for filename in os.listdir("outputs"):
        if filename.endswith(".txt"):
            with open(f"outputs/{filename}", "r") as f:
                text = f.read()

            # Get AI analysis
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[{
                    "role": "system",
                    "content": """Create a tender report with:
                    1. Executive Summary (3 sentences)
                    2. Key Requirements (bulleted list)
                    3. Technical Specifications (table format)
                    4. Risk Analysis (2 columns: Risk/Mitigation)"""
                }, {
                    "role": "user",
                    "content": text[:15000]
                }]
            )
            
            # Create document
            doc = Document()
            
            # Custom styles
            styles = doc.styles
            styles['Heading 1'].font.color.rgb = RGBColor(0, 32, 96)
            
            # Add elements
            create_cover_page(doc)
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(response.choices[0].message.content.split('\n')[0])
            
            add_methodology_table(doc)
            
            # Save
            doc.save(f"outputs/{filename.replace('.txt', '_report.docx')}")

if __name__ == "__main__":
    generate_reports()
